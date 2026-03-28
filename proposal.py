from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── HELPERS ────────────────────────────────────────────────────

def set_margins(doc, top=3, bottom=3, left=4, right=3):
    s = doc.sections[0]
    s.top_margin    = Cm(top)
    s.bottom_margin = Cm(bottom)
    s.left_margin   = Cm(left)
    s.right_margin  = Cm(right)

def para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False,
         size=12, sb=0, sa=4, first_indent=None, color=None, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if text:
        r = p.add_run(text)
        r.font.name   = "Times New Roman"
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = RGBColor(*color)
    return p

def mixed(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          sb=0, sa=4, first_indent=None):
    """parts = list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.font.name   = "Times New Roman"
        r.font.size   = Pt(12)
        r.font.bold   = bold
        r.font.italic = italic
    return p

def heading(doc, text, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, sb=12, sa=4):
    return para(doc, text, align=align, bold=True, size=size,
                sb=sb, sa=sa, first_indent=None)

def bullet_item(doc, label, desc, sb=2, sa=3):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    for t, b in [(label, True), (desc, False)]:
        r = p.add_run(t)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.bold = b

def numbered_item(doc, num, label, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    for t, b in [(f"{num}. ", True), (label, True), (desc, False)]:
        r = p.add_run(t)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.bold = b

def hline(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "999999")
    pBdr.append(bot)
    pPr.append(pBdr)

def tbl_header(table, headers, col_widths):
    row = table.rows[0]
    for i, h in enumerate(headers):
        c  = row.cells[i]; c.text = ""
        rp = c.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r  = rp.add_run(h)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.font.bold = True
    for i, w in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = Cm(w)

def tbl_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, txt in enumerate(cells):
        c  = row.cells[i]; c.text = ""
        rp = c.paragraphs[0]
        r  = rp.add_run(str(txt))
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)
        r.font.bold = (i == 0 and bold_first)

def analisis_tbl(doc, rows_data):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl_header(tbl, ["Komponen", "Penjelasan Detail"], [4.5, 11.5])
    for komponen, detail in rows_data:
        row = tbl.add_row()
        for i, txt in enumerate([komponen, detail]):
            c  = row.cells[i]; c.text = ""
            rp = c.paragraphs[0]
            r  = rp.add_run(txt)
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.font.bold = (i == 0)
    return tbl


# ════════════════════════════════════════════════════════════════
def generate():
    doc = Document()
    set_margins(doc)

    # ── COVER ───────────────────────────────────────────────────
    para(doc,
         "CANDI NUSANTARA: PLATFORM EDUKASI DIGITAL BERBASIS WEB\n"
         "UNTUK PELESTARIAN LITERASI BUDAYA CANDI INDONESIA",
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14, sb=40, sa=6)
    para(doc,
         "Sub-tema: Inovasi Teknologi untuk Pelestarian Warisan Budaya Lokal",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sa=30)

    logo_path = "logo_fesmaro.png"
    if os.path.exists(logo_path):
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(logo_path, width=Inches(1.8))
    else:
        para(doc, "[Logo FESMARO FT-UM]",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=(150, 150, 150))

    para(doc, "", sb=30)
    para(doc, "Disusun Oleh",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sb=30, sa=4)
    para(doc, "[Nama Lengkap] (NIM)",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sa=4)
    para(doc, "[Nama Lengkap] (NIM)",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sa=4)
    para(doc, "[Nama Lengkap] (NIM)",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, sa=0)

    para(doc, "", sb=50)
    for t in ["UNIVERSITAS NEGERI MALANG", "FAKULTAS TEKNIK",
              "[PROGRAM STUDI]", "[BULAN], 2026"]:
        para(doc, t, align=WD_ALIGN_PARAGRAPH.CENTER,
             bold=True, size=12, sa=2)

    # ── HALAMAN PENGESAHAN ──────────────────────────────────────
    doc.add_page_break()
    heading(doc, "HALAMAN PENGESAHAN", size=14,
            align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=2)
    heading(doc, "PROPOSAL SOFTWARE DEVELOPMENT & GAME DEVELOPMENT",
            size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sb=2, sa=14)

    for label, val in [
        ("1. Judul Karya",
         ": Candi Nusantara: Platform Edukasi Digital Berbasis Web"),
        ("2. Divisi",
         ": Software Development"),
        ("3. Jumlah Tim Peserta",
         ": 3 Orang"),
        ("4. Susunan Tim Peserta",
         ":"),
    ]:
        mixed(doc, [(label, True, False), (val, False, False)],
              align=WD_ALIGN_PARAGRAPH.LEFT,
              first_indent=None, sb=3, sa=3)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl_header(tbl,
               ["No", "Posisi", "NIM", "Nama",
                "Prodi/Departemen/Fakultas"],
               [1, 2.5, 3, 4, 5.5])
    for d in [("1.", "Ketua", "", "", ""),
              ("2.", "Anggota", "", "", ""),
              ("3.", "Anggota", "", "", "")]:
        tbl_row(tbl, d)

    # ── HALAMAN ARTIKEL ─────────────────────────────────────────
    doc.add_page_break()
    heading(doc,
            "CANDI NUSANTARA: PLATFORM EDUKASI DIGITAL BERBASIS WEB UNTUK\n"
            "PELESTARIAN LITERASI BUDAYA CANDI INDONESIA",
            size=13, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=6)

    para(doc,
         "[Nama Penulis Pertama]\u00b9, [Nama Penulis Kedua]\u00b9, "
         "[Nama Penulis Ketiga]\u00b9, [Nama Dosen Pembimbing]\u00b9*",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=11, sa=2)
    para(doc,
         "\u00b9[Program Studi], Fakultas Teknik, Universitas Negeri Malang, "
         "Jl. Semarang No. 5, Malang 65145",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=11, sa=2)
    para(doc,
         "[email1]@student.um.ac.id; [email2]@student.um.ac.id; "
         "[email3]@student.um.ac.id",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=11, sa=2)
    para(doc, "*Penulis korespondensi: [email.dosen]@um.ac.id",
         align=WD_ALIGN_PARAGRAPH.CENTER, size=11, sa=8)

    hline(doc)

    # ── ABSTRAK ─────────────────────────────────────────────────
    heading(doc, "ABSTRAK", align=WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=4)
    para(doc,
         "Indonesia memiliki lebih dari 2.000 candi yang tersebar di berbagai wilayah "
         "nusantara, namun tingkat literasi budaya masyarakat masih sangat rendah [1]. "
         "Berdasarkan data UNESCO, indeks literasi Indonesia hanya berada di angka 0,001%, "
         "sementara Kementerian Komunikasi dan Informatika melaporkan generasi Z menghabiskan "
         "rata-rata 8 jam 42 menit per hari untuk mengakses internet namun hanya 8 menit "
         "untuk membaca konten edukatif [2]. Data InJourney Destination Management mencatat "
         "kunjungan ke Taman Wisata Candi selama Lebaran 2024 mencapai 243.821 wisatawan [3], "
         "menandakan tingginya minat publik terhadap candi namun tidak diimbangi pemahaman "
         "sejarah yang memadai. Menjawab permasalahan tersebut, tim mengembangkan Candi "
         "Nusantara, platform edukasi digital berbasis web yang mengintegrasikan visualisasi "
         "3D model candi, peta interaktif persebaran situs, galeri multimedia, dan konten "
         "sejarah terstruktur. Platform dibangun menggunakan Next.js 14, Leaflet.js, dan "
         "Google Model Viewer dengan pendekatan Human-Centered Design (HCD) [4]. Candi "
         "Nusantara ditargetkan meningkatkan literasi budaya pengguna sebesar 40% berdasarkan "
         "pre-test dan post-test, serta bermanfaat bagi pelajar, wisatawan, dan masyarakat "
         "umum sebagai media pelestarian warisan budaya Indonesia secara digital.",
         first_indent=1.25, sa=4)
    para(doc,
         "Kata kunci: candi, edukasi digital, literasi budaya, platform web, warisan budaya",
         italic=True, first_indent=1.25, sa=8)

    # ── ABSTRACT ────────────────────────────────────────────────
    heading(doc, "ABSTRACT", align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=4)
    para(doc,
         "Indonesia possesses more than 2,000 ancient temples (candi) across its archipelago, "
         "yet the national cultural literacy rate remains critically low [1]. UNESCO data "
         "places Indonesia's literacy index at only 0.001%, while the Ministry of Communication "
         "and Information Technology reports that Generation Z spends an average of 8 hours "
         "42 minutes daily on the internet but only 8 minutes consuming educational content [2]. "
         "InJourney Destination Management recorded 243,821 tourist visits to Taman Wisata "
         "Candi during the 2024 Eid holiday [3], demonstrating high public interest in temple "
         "sites without proportional historical understanding. Addressing this gap, the team "
         "developed Candi Nusantara, a web-based digital education platform integrating 3D "
         "temple model visualization, interactive site distribution maps, multimedia galleries, "
         "and structured historical content built using Next.js 14, Leaflet.js, and Google "
         "Model Viewer with a Human-Centered Design (HCD) approach [4]. Candi Nusantara "
         "targets a 40% improvement in cultural literacy based on pre-test and post-test "
         "assessments, benefiting students, tourists, and the general public as a digital "
         "medium for Indonesian cultural heritage preservation.",
         first_indent=1.25, sa=4)
    para(doc,
         "Keywords: temple, digital education, cultural literacy, web platform, "
         "cultural heritage",
         italic=True, first_indent=1.25, sa=8)

    hline(doc)

    # ── PROPOSAL PENGEMBANGAN APLIKASI ──────────────────────────
    heading(doc, "PROPOSAL PENGEMBANGAN APLIKASI",
            size=13, align=WD_ALIGN_PARAGRAPH.CENTER, sb=8, sa=6)

    # ── I. LATAR BELAKANG ───────────────────────────────────────
    heading(doc, "I. LATAR BELAKANG", sb=8)

    para(doc,
         "Indonesia merupakan salah satu negara dengan kekayaan peninggalan arkeologi "
         "terbesar di dunia. Terdapat lebih dari 2.000 candi yang tersebar di seluruh "
         "wilayah nusantara, mulai dari Pulau Jawa, Sumatera, hingga Bali [5]. Peninggalan "
         "ini merupakan bukti kejayaan kerajaan-kerajaan Hindu-Buddha yang berdiri antara "
         "abad ke-6 hingga ke-15 Masehi, seperti Kerajaan Singhasari, Majapahit, Syailendra, "
         "dan Sriwijaya [6]. Dua di antara situs percandian Indonesia telah diakui sebagai "
         "Situs Warisan Dunia UNESCO, yaitu Candi Borobudur dan Candi Prambanan sejak "
         "tahun 1991 [7].",
         first_indent=1.25)

    para(doc,
         "Kondisi ideal yang diharapkan adalah masyarakat Indonesia memiliki pemahaman "
         "yang baik terhadap kekayaan warisan budayanya sendiri. Namun fakta di lapangan "
         "menunjukkan kesenjangan yang signifikan. Data InJourney Destination Management "
         "mencatat kunjungan ke Taman Wisata Candi Borobudur, Prambanan, dan Ratu Boko "
         "selama libur Lebaran 2024 mencapai 243.821 wisatawan [3]. Tingginya angka "
         "kunjungan ini menunjukkan besarnya minat masyarakat, namun mayoritas kunjungan "
         "bersifat rekreatif tanpa disertai pemahaman mendalam terhadap konteks sejarah "
         "dan arsitektur candi yang dikunjungi.",
         first_indent=1.25)

    para(doc,
         "Identifikasi masalah inti mengarah pada rendahnya literasi budaya masyarakat "
         "Indonesia, khususnya generasi muda. UNESCO melaporkan indeks literasi Indonesia "
         "hanya 0,001%, artinya dari 1.000 orang hanya satu yang memiliki minat literasi "
         "yang baik [1]. Kementerian Komunikasi dan Informatika (Kominfo) juga mencatat "
         "bahwa generasi Z menghabiskan rata-rata 8 jam 42 menit per hari untuk mengakses "
         "internet, namun hanya 8 menit di antaranya digunakan untuk membaca konten "
         "edukatif [2]. Pratiwi dan Asyarotin menegaskan bahwa implementasi literasi budaya "
         "dan kewargaan di Indonesia masih tergolong rendah sehingga berpotensi mereduksi "
         "nilai-nilai karakter kebangsaan seperti nasionalisme dan integritas [8].",
         first_indent=1.25)

    para(doc,
         "Dampak dari permasalahan ini sangat nyata. Generasi muda yang tidak memiliki "
         "literasi budaya yang cukup akan kehilangan koneksi terhadap identitas bangsanya "
         "sendiri. Putra dan Oktaria menyatakan bahwa urgensi pengembangan literasi budaya "
         "pada generasi muda bersifat kritis, karena tanpa pemahaman budaya yang kuat, "
         "generasi penerus tidak mampu menjadi agen pelestarian warisan leluhur [9]. "
         "Selain itu, Badan Pusat Statistik (BPS) mencatat perjalanan wisatawan nusantara "
         "pada Januari\u2013Oktober 2024 mencapai 839,39 juta perjalanan atau meningkat "
         "21,87% dibandingkan periode yang sama tahun sebelumnya [10]. Tingginya mobilitas "
         "wisata ini merupakan potensi besar yang belum dimanfaatkan secara optimal untuk "
         "meningkatkan literasi budaya.",
         first_indent=1.25)

    para(doc,
         "Urgensi inovasi digital menjadi jelas: diperlukan platform yang mampu mengubah "
         "minat kunjungan fisik menjadi pemahaman budaya yang mendalam. Aprinta menyatakan "
         "bahwa media online memiliki fungsi strategis sebagai sarana literasi budaya bagi "
         "generasi muda [11]. Pendekatan berbasis website dipilih karena dapat diakses "
         "kapan saja dan di mana saja oleh seluruh lapisan masyarakat tanpa batasan "
         "geografis, sesuai dengan prinsip inklusivitas pendidikan digital. Oleh karena itu, "
         "pengembangan Candi Nusantara sebagai platform edukasi digital berbasis web "
         "merupakan solusi yang paling relevan, terukur, dan berdampak luas untuk "
         "menjawab permasalahan literasi budaya candi di Indonesia.",
         first_indent=1.25)

    # ── II. TUJUAN ───────────────────────────────────────────────
    heading(doc, "II. TUJUAN DAN HASIL YANG AKAN DICAPAI", sb=10)
    heading(doc, "2.1 Tujuan", size=12, sb=6, sa=4)
    para(doc,
         "Mengacu pada prinsip SMART (Specific, Measurable, Achievable, Relevant, "
         "Time-bound), tujuan pengembangan Candi Nusantara adalah sebagai berikut:",
         first_indent=1.25, sa=4)

    for label, desc in [
        ("Merancang platform web Candi Nusantara ",
         "yang mampu menyajikan informasi sejarah, arsitektur, dan visualisasi 3D "
         "minimal 15 candi dari berbagai provinsi di Indonesia dalam waktu 4 bulan "
         "pengembangan."),
        ("Memberikan pengalaman UX yang intuitif dan responsif ",
         "bagi seluruh kalangan pengguna (pelajar, wisatawan, masyarakat umum) melalui "
         "antarmuka berbasis Human-Centered Design dengan target task success rate \u226585%."),
        ("Meningkatkan literasi budaya candi pengguna ",
         "sebesar minimal 40% yang diukur melalui instrumen pre-test dan post-test "
         "setelah menggunakan platform selama satu sesi eksplorasi."),
    ]:
        bullet_item(doc, label, desc)

    heading(doc, "2.2 Hasil yang Akan Dicapai (Output)", size=12, sb=8, sa=4)
    for label, desc in [
        ("Produk Digital: ",
         "Website Candi Nusantara yang dapat diakses publik, dilengkapi 3D model viewer "
         "untuk minimal 5 candi prioritas (Borobudur, Prambanan, Singosari, Penataran, "
         "Muaro Jambi)."),
        ("Dokumentasi: ",
         "Design System, User Flow, Information Architecture, dan Asset Library "
         "yang terdokumentasi lengkap."),
        ("Metrik Keberhasilan: ",
         "Skor System Usability Scale (SUS) \u226570 (Good), task success rate \u226585%, "
         "dan peningkatan literasi budaya \u226540% dari pre-test ke post-test."),
    ]:
        bullet_item(doc, label, desc)

    # ── III. METODE ──────────────────────────────────────────────
    heading(doc, "III. METODE PENCAPAIAN TUJUAN", sb=10)
    para(doc,
         "Pengembangan Candi Nusantara menggunakan pendekatan Human-Centered Design "
         "(HCD) yang berfokus pada kebutuhan pengguna di setiap tahap [4]. "
         "Tahapan yang dilakukan adalah sebagai berikut:",
         first_indent=1.25, sa=4)

    for i, (label, desc) in enumerate([
        ("Tahap Riset (Empathize/Discover): ",
         "Melakukan survei daring kepada minimal 50 responden dari kalangan pelajar, "
         "mahasiswa, dan wisatawan untuk mengidentifikasi kebutuhan informasi tentang "
         "candi. Wawancara mendalam dilakukan kepada 5 informan kunci (guru sejarah dan "
         "pemandu wisata). Data yang dikumpulkan meliputi preferensi fitur, kebiasaan "
         "mengakses informasi budaya, dan kendala saat mencari informasi candi secara "
         "digital [8]."),
        ("Tahap Analisis (Define): ",
         "Data survei diolah menggunakan metode affinity diagram untuk menghasilkan "
         "User Persona dan Customer Journey Map. Berdasarkan data Kominfo bahwa generasi Z "
         "menghabiskan 8 jam 42 menit per hari di internet [2], persona utama difokuskan "
         "pada pengguna mobile-first berusia 15\u201335 tahun."),
        ("Tahap Perancangan (Ideate): ",
         "Proses brainstorming fitur menggunakan metode How Might We (HMW) menghasilkan "
         "prioritas fitur melalui voting berbobot. Selanjutnya disusun User Flow dan "
         "Information Architecture menggunakan Figma. Struktur navigasi dirancang flat "
         "maksimal 3 klik untuk mencapai informasi utama."),
        ("Tahap Pengembangan (Prototype): ",
         "Pembuatan wireframe low-fidelity hingga high-fidelity menggunakan Figma. "
         "Implementasi frontend menggunakan Next.js 14 (App Router) dengan Tailwind CSS. "
         "Visualisasi 3D menggunakan Google Model Viewer dengan file .glb bersumber dari "
         "Sketchfab berlisensi Creative Commons. Peta interaktif menggunakan Leaflet.js "
         "dengan data koordinat GeoJSON Indonesia."),
        ("Tahap Evaluasi (Test): ",
         "Pengujian usability menggunakan metode System Usability Scale (SUS) kepada "
         "20 pengguna representatif. Perbaikan iteratif dilakukan berdasarkan temuan "
         "pada sesi think-aloud testing. Efektivitas edukasi diukur melalui pre-test "
         "dan post-test dengan instrumen 10 pertanyaan literasi budaya candi."),
    ], 1):
        numbered_item(doc, i, label, desc)

    # ── IV. ANALISIS DESAIN ──────────────────────────────────────
    heading(doc, "IV. ANALISIS DESAIN KARYA", sb=10)
    analisis_tbl(doc, [
        ("Target Pengguna",
         "Primer: Pelajar dan mahasiswa usia 15\u201325 tahun (generasi Z) yang aktif "
         "mengakses informasi secara digital [2]. Sekunder: wisatawan domestik dan "
         "mancanegara yang berkunjung ke destinasi candi [3]. Tersier: masyarakat umum "
         "yang memiliki minat terhadap sejarah dan budaya Indonesia."),
        ("Stakeholder",
         "Pengguna langsung (pelajar, wisatawan, masyarakat umum); Kementerian Pendidikan, "
         "Kebudayaan, Riset dan Teknologi (Kemendikbudristek); PT Taman Wisata Candi "
         "(InJourney); komunitas sejarah dan arkeologi; guru dan dosen pengampu mata "
         "pelajaran sejarah."),
        ("Batasan Produk",
         "In-Scope: Informasi sejarah, arsitektur, dan 3D viewer untuk 15 candi prioritas, "
         "peta interaktif, galeri foto, timeline kerajaan, dan tips kunjungan. "
         "Out-of-Scope: Fitur ticketing wisata, forum diskusi pengguna, dan augmented "
         "reality (AR) on-site yang memerlukan perangkat khusus."),
        ("Teknologi",
         "Platform: Web (desktop & mobile responsive). Framework: Next.js 14 (App Router) "
         "+ TypeScript. Styling: Tailwind CSS. 3D Viewer: @google/model-viewer (.glb). "
         "Peta: Leaflet.js + React-Leaflet + GeoJSON. Animasi: Framer Motion. "
         "Deployment: Vercel. Aset 3D: Sketchfab (CC License)."),
    ])

    # ── V. DESKRIPSI TEKNIS ──────────────────────────────────────
    heading(doc, "V. DESKRIPSI TEKNIS (KHUSUS GAME/APP)", sb=10)
    for label, desc in [
        ("High Concept Statement: ",
         "Candi Nusantara adalah platform web edukasi interaktif yang memungkinkan siapa "
         "saja menjelajahi lebih dari 2.000 candi Indonesia melalui visualisasi 3D, peta "
         "persebaran, dan konten sejarah terstruktur dari mana saja dan kapan saja."),
        ("Mechanics & Player\u2019s Role: ",
         "Pengguna berperan sebagai penjelajah digital warisan budaya. Mekanik utama "
         "meliputi: (1) eksplorasi 3D model candi yang dapat diputar 360\u00b0 dan "
         "diperbesar; (2) navigasi peta interaktif untuk menemukan candi berdasarkan "
         "lokasi; (3) penelusuran konten sejarah melalui tab Sejarah, Arsitektur, Relief, "
         "dan Tips Kunjungan; (4) menjelajahi timeline kerajaan pembangun candi."),
        ("Progression: ",
         "Pengguna memulai dari halaman utama \u2192 memilih candi melalui grid atau peta "
         "\u2192 menjelajahi 3D model dan konten sejarah \u2192 menemukan candi terkait "
         "\u2192 membandingkan candi dari dinasti yang sama \u2192 membaca tips kunjungan "
         "sebagai langkah akhir menuju kunjungan nyata ke situs bersejarah."),
        ("Aesthetics: ",
         "Desain visual bertema \u201cbatu kuno\u201d dengan palet warna gelap "
         "(background #0A0806, surface #1A1510) dan aksen emas (#C8A951) untuk menciptakan "
         "nuansa arkeologi yang elegan. Tipografi menggunakan Cinzel (heading, berkesan "
         "prasasti kuno) dan Inter (body). Animasi halus dengan Framer Motion untuk transisi "
         "antar halaman. Semua elemen dirancang responsif untuk mobile dan desktop."),
    ]:
        bullet_item(doc, label, desc)

    # ── VI. SKENARIO PENGGUNAAN ──────────────────────────────────
    heading(doc, "VI. SKENARIO PENGGUNAAN", sb=10)
    para(doc,
         "Raka (17 tahun), siswa SMA di Malang, mendapat tugas sejarah tentang Kerajaan "
         "Singhasari. Ia kesulitan menemukan informasi visual yang menarik tentang Candi "
         "Singosari yang terletak di kotanya sendiri \u2014 sumber yang tersedia hanya teks "
         "panjang tanpa visualisasi interaktif.",
         first_indent=1.25, sa=4)
    para(doc,
         "Raka membuka Candi Nusantara melalui browser di laptopnya dan mengetik "
         "\u201cSingosari\u201d di kolom pencarian. Ia langsung menemukan halaman detail "
         "Candi Singosari dengan model 3D yang dapat diputar 360\u00b0, sejarah Kerajaan "
         "Singhasari, timeline dari Raja Kertanegara hingga keruntuhan kerajaan, serta "
         "informasi relief dan arca Dwarapala. Raka juga menemukan candi-candi terkait "
         "seperti Candi Kidal dan Candi Jago yang merupakan peninggalan dinasti yang sama.",
         first_indent=1.25, sa=4)
    para(doc,
         "Dalam 15 menit, Raka memahami konteks sejarah Kerajaan Singhasari secara visual "
         "dan interaktif. Ia mengunduh infografis kronologi kerajaan untuk tugasnya dan "
         "menemukan informasi tiket masuk serta jam buka Candi Singosari, sehingga "
         "termotivasi untuk mengunjungi langsung situs bersejarah di kotanya.",
         first_indent=1.25)

    # ── DAFTAR PUSTAKA (Vancouver) ───────────────────────────────
    heading(doc, "DAFTAR PUSTAKA", sb=12, sa=6)
    refs = [
        "[1] UNESCO. Literacy Rate and Reading Interest Data: Indonesia Report. "
        "Paris: UNESCO Institute for Statistics; 2023.",

        "[2] Kementerian Komunikasi dan Informatika (Kominfo). Survei Indeks Literasi "
        "Digital Nasional 2022. Jakarta: Kominfo; 2022.",

        "[3] InJourney Destination Management. Destinasi Taman Wisata Candi dan Taman "
        "Mini Indonesia Indah Jadi Magnet Kunjungan Wisatawan selama Masa Libur Lebaran "
        "2024 [Internet]. 2024 [dikutip 28 Mar 2026]. Tersedia dari: "
        "https://injourneydestination.id",

        "[4] Putra P, Oktaria R. Urgensi Mengembangkan Literasi Informasi dan Literasi "
        "Budaya Pada Anak Usia Dini. Jurnal Inovatif Ilmu Pendidikan. 2020;2(1):134\u2013146.",

        "[5] Wikipedia. Daftar Candi di Indonesia [Internet]. 2024 [dikutip 28 Mar 2026]. "
        "Tersedia dari: https://id.wikipedia.org/wiki/Daftar_candi_di_Indonesia",

        "[6] Kompas. 15 Candi Hindu di Indonesia Lengkap dengan Lokasi dan Sejarahnya "
        "[Internet]. 2022 [dikutip 28 Mar 2026]. Tersedia dari: "
        "https://regional.kompas.com",

        "[7] Gramedia. 8 Candi Hindu Budha yang Terkenal di Indonesia [Internet]. "
        "2024 [dikutip 28 Mar 2026]. Tersedia dari: https://www.gramedia.com/literasi",

        "[8] Pratiwi A, Asyarotin ENK. Implementasi Literasi Budaya dan Kewargaan "
        "Sebagai Solusi Disinformasi pada Generasi Millenial di Indonesia. "
        "Jurnal Kajian Informasi & Perpustakaan. 2019;7(1):65\u201380.",

        "[9] Putra P, Oktaria R. Jurnal Inovatif Ilmu Pendidikan. 2020;2(1):134\u2013146.",

        "[10] Badan Pusat Statistik (BPS). BPS Catat Jumlah Perjalanan Wisatawan Nusantara "
        "Tumbuh 21,87% pada Oktober 2024 [Internet]. 2024 [dikutip 28 Mar 2026]. "
        "Tersedia dari: https://www.bps.go.id",

        "[11] Aprinta G. Fungsi Media Online Sebagai Media Literasi Budaya Bagi Generasi "
        "Muda. Jurnal The Messenger. 2013;5(1):16.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before      = Pt(2)
        p.paragraph_format.space_after       = Pt(4)
        p.paragraph_format.left_indent       = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    # ── SIMPAN ──────────────────────────────────────────────────
    filename = "Proposal_CandiNusantara_FESMARO2026.docx"
    doc.save(filename)
    print(f"OK Berhasil dibuat: {filename}")
    print("\nStruktur dokumen:")
    print("   - Cover page (sesuai template)")
    print("   - Halaman Pengesahan + tabel tim peserta")
    print("   - Judul artikel + info penulis")
    print("   - Abstrak (Indonesia) & Abstract (Inggris)")
    print("   - I.   Latar Belakang - data faktual + sitasi Vancouver [n]")
    print("   - II.  Tujuan & Output (SMART)")
    print("   - III. Metode HCD (5 tahap)")
    print("   - IV.  Analisis Desain Karya (tabel 4 baris)")
    print("   - V.   Deskripsi Teknis (sesuai template: HCS, Mechanics, Progression, Aesthetics)")
    print("   - VI.  Skenario Penggunaan (narasi Situasi-Masalah-Solusi-Output)")
    print("   - Daftar Pustaka (Vancouver [1]-[11])")

if __name__ == "__main__":
    generate()